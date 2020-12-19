import docx
from docx2pdf import convert
doc=docx.Document('testing.docx')

company_name="Bio matrix"
for pra,i in zip(doc.paragraphs,range(0,len(doc.paragraphs))):
    print(i," => ",pra.text)

# clear the company name
doc.paragraphs[4]._p.clear()

#adding the company name in the document
doc.paragraphs[4].add_run('\t\t\t'+company_name).bold=True

#subject
subject="\t SUBJECT: " \
        " AGREEMENT BETWEEN " \
        + company_name.capitalize()  +  \
        " \n\t\t      AND RECRUITMENT HR EXPERTS (Division of RHE Career Movers Pvt Ltd)."


#clear the subject parapgraph
doc.paragraphs[7]._p.clear()

#adding subject to the document.
doc.paragraphs[7].add_run(subject).bold=True

#checking rate
for run,i in zip(doc.paragraphs[12].runs,range(0,len(doc.paragraphs[12].runs))):
    print(i, " =>" ,run.text)


#print("older ->"+doc.paragraphs[12].runs[2].text)
doc.paragraphs[12].runs[2].text="8.33";
#print("newer ->"+doc.paragraphs[12].runs[2].text)

#for date
for run,i in zip(doc.paragraphs[2].runs,range(0,len(doc.paragraphs[2].runs))):
    print(i, " =>" ,run.text)
doc.paragraphs[2].runs[11].text='9th octomber 2020'

#for replacement period
for run,i in zip(doc.paragraphs[16].runs,range(0,len(doc.paragraphs[16].runs))):
    print(i, " =>" ,run.text)
doc.paragraphs[16].runs[3].text="90 days"

#for payment terms
for run,i in zip(doc.paragraphs[18].runs,range(0,len(doc.paragraphs[18].runs))):
    print(i, " =>" ,run.text)
doc.paragraphs[18].runs[6].text="30 days"

#for Validity
for run,i in zip(doc.paragraphs[32].runs,range(0,len(doc.paragraphs[32].runs))):
    print(i, " =>" ,run.text)

doc.paragraphs[32].runs[5].text="180 days"
doc.paragraphs[32].runs[5].bold=True



#For company in last
for run,i in zip(doc.paragraphs[40].runs,range(0,len(doc.paragraphs[40].runs))):
    print(i, " =>" ,run.text)
doc.paragraphs[40].runs[14].text=company_name.capitalize()


#save the document
doc.save(company_name+'.docx')
