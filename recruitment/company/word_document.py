import docx

def generate_document(dict):
    doc=docx.Document('testing.docx')

    company_name=dict["company_name"]
    rate=dict["rate"]
    date = dict["date"]
    payment_terms=dict["payment_terms"]
    validity = dict["validity"]
    replacement_period=dict["replacement_period"]

    # clear the company name
    doc.paragraphs[4]._p.clear()

    #adding the company name in the document
    doc.paragraphs[4].add_run('\t\t\t'+company_name).bold=True

    #subject
    subject="\t SUBJECT: " \
            " AGREEMENT BETWEEN " \
            + company_name +  \
            " \n\t\t      AND RECRUITMENT HR EXPERTS (Division of RHE Career Movers Pvt Ltd)."


    #clear the subject parapgraph
    doc.paragraphs[7]._p.clear()

    #adding subject to the document.
    doc.paragraphs[7].add_run(subject).bold=True

    #checking rate
    #print("older ->"+doc.paragraphs[12].runs[2].text)
    doc.paragraphs[12].runs[2].text=str(rate);
    #print("newer ->"+doc.paragraphs[12].runs[2].text)

    #for date
    doc.paragraphs[2].runs[11].text=str(date)

    #for replacement period
    doc.paragraphs[16].runs[3].text=replacement_period

    #for payment terms
    doc.paragraphs[18].runs[6].text=payment_terms

    #for Validity
    doc.paragraphs[32].runs[5].text=validity
    doc.paragraphs[32].runs[5].bold=True

    #For company in last
    doc.paragraphs[40].runs[14].text=company_name

    #save the document
    doc.save(company_name+'.docx')

    return company_name+'.docx'